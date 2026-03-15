#!/usr/bin/env python3
"""
生成 Neuroscience Daily Digest Word 文档
使用简约高级风格排版
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from datetime import datetime
import re

class WordGenerator:
    """Word文档生成器"""
    
    # 配色方案
    COLORS = {
        'primary': (26, 54, 93),      # #1a365d 深蓝
        'secondary': (44, 82, 130),   # #2c5282 中蓝
        'text': (45, 55, 72),         # #2d3748 深灰
        'muted': (74, 85, 104),       # #4a5568 中灰
        'light': (113, 128, 150),     # #718096 浅灰
        'border': (226, 232, 240),    # #e2e8f0 边框灰
    }
    
    def __init__(self):
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """设置文档基础格式"""
        # 页面设置
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)
    
    def _set_run_style(self, run, font_name='Microsoft YaHei', font_size=11, 
                       bold=False, italic=False, color=None):
        """设置文本运行样式"""
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def add_title_page(self, date_str, count):
        """添加标题页"""
        # 主标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('🧠 Neuroscience Daily Digest')
        self._set_run_style(run, font_size=22, bold=True, 
                           color=self.COLORS['primary'])
        
        # 副标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{date_str}')
        self._set_run_style(run, font_size=14, 
                           color=self.COLORS['light'])
        
        # 统计信息
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'精选 {count} 篇文献')
        self._set_run_style(run, font_size=12, 
                           color=self.COLORS['muted'])
        
        # 分隔空行
        self.doc.add_paragraph()
        self.doc.add_paragraph()
    
    def add_article(self, number, en_title, cn_title, abstract, journal, date):
        """添加一篇文章"""
        # 分隔线效果（空段落）
        self.doc.add_paragraph()
        
        # 序号 + 英文标题
        p = self.doc.add_paragraph()
        run = p.add_run(f"{number}. {en_title}")
        self._set_run_style(run, font_size=14, bold=True, 
                           color=self.COLORS['primary'])
        p.paragraph_format.space_after = Pt(6)
        
        # 中文标题
        p = self.doc.add_paragraph()
        run = p.add_run(cn_title)
        self._set_run_style(run, font_size=12, italic=True, 
                           color=self.COLORS['muted'])
        p.paragraph_format.space_after = Pt(6)
        
        # 摘要
        p = self.doc.add_paragraph()
        run = p.add_run(f"摘要：{abstract}")
        self._set_run_style(run, font_size=11, 
                           color=self.COLORS['text'])
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        
        # 来源
        p = self.doc.add_paragraph()
        run = p.add_run(f"📄 {journal} | {date}")
        self._set_run_style(run, font_size=10, italic=True, 
                           color=self.COLORS['light'])
        p.paragraph_format.space_after = Pt(12)
    
    def add_statistics(self, count, sources):
        """添加统计信息"""
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('--- 今日统计 ---')
        self._set_run_style(run, font_size=12, bold=True, 
                           color=self.COLORS['secondary'])
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'总文章数：{count}篇')
        self._set_run_style(run, font_size=11, 
                           color=self.COLORS['text'])
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'来源：{sources}')
        self._set_run_style(run, font_size=10, 
                           color=self.COLORS['light'])
    
    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
        print(f"✅ Word文档已生成: {output_path}")


def parse_markdown(md_content):
    """解析Markdown内容"""
    articles = []
    lines = md_content.split('\n')
    
    current_article = {}
    in_abstract = False
    
    for line in lines:
        line = line.strip()
        
        # 匹配文章标题行: ### 1. Title
        if line.startswith('### ') and '. ' in line:
            if current_article and 'title' in current_article:
                articles.append(current_article)
            current_article = {}
            parts = line.replace('### ', '').split('. ', 1)
            current_article['number'] = parts[0]
            current_article['title'] = parts[1]
            in_abstract = False
        
        # 匹配中文标题: *中文标题*
        elif line.startswith('*') and line.endswith('*') and 'title' in current_article:
            current_article['cn_title'] = line.strip('*')
        
        # 匹配摘要
        elif line.startswith('**摘要**：'):
            current_article['abstract'] = line.replace('**摘要**：', '')
            in_abstract = True
        
        # 继续摘要（多行）
        elif in_abstract and line and not line.startswith('📄'):
            current_article['abstract'] += line
        
        # 匹配来源: 📄 Journal | Date
        elif line.startswith('📄'):
            parts = line.replace('📄 ', '').split(' | ')
            if len(parts) == 2:
                current_article['journal'] = parts[0].strip()
                current_article['date'] = parts[1].strip()
            in_abstract = False
    
    if current_article and 'title' in current_article:
        articles.append(current_article)
    
    return articles


def main():
    """主函数"""
    import sys
    
    if len(sys.argv) < 3:
        print("用法: python generate_word.py --input input.md --output output.docx")
        return
    
    input_file = sys.argv[2]
    output_file = sys.argv[4] if len(sys.argv) > 4 else 'neuro_digest.docx'
    
    # 读取Markdown
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 解析文章
    articles = parse_markdown(md_content)
    
    # 生成Word
    gen = WordGenerator()
    
    # 提取日期（从文件名或内容）
    date_str = datetime.now().strftime('%Y年%m月%d日')
    
    # 添加标题页
    gen.add_title_page(date_str, len(articles))
    
    # 添加文章
    for article in articles:
        gen.add_article(
            number=article.get('number', '1'),
            en_title=article.get('title', ''),
            cn_title=article.get('cn_title', ''),
            abstract=article.get('abstract', ''),
            journal=article.get('journal', 'Unknown'),
            date=article.get('date', date_str)
        )
    
    # 添加统计
    sources = ', '.join(set(a.get('journal', '') for a in articles[:5]))
    gen.add_statistics(len(articles), sources)
    
    # 保存
    gen.save(output_file)


if __name__ == '__main__':
    main()
